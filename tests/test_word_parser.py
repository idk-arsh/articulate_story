"""
Test Word Parser
"""

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt
from app.parsers.word_parser import WordParser


def create_sample_word_file():
    """Create a sample Storyline-style Word translation file"""
    
    doc = Document()
    
    # Add title
    doc.add_heading('Storyline Translation Export', 0)
    
    # Create translation table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ID'
    header_cells[1].text = 'Original Text'
    header_cells[2].text = 'Translation'
    
    # Add sample content
    sample_data = [
        ('slide1_text1', 'Welcome to the course!', ''),
        ('slide1_text2', 'Click the button to continue.', ''),
        ('slide2_text1', 'You have %TotalQuestions%% questions.', ''),
        ('slide2_text2', 'Press <x id="1"/> to start.', ''),
        ('slide3_text1', 'Thank you for completing!', ''),
    ]
    
    for id_text, original, translation in sample_data:
        row_cells = table.add_row().cells
        row_cells[0].text = id_text
        row_cells[1].text = original
        row_cells[2].text = translation
    
    # Save
    output_path = Path(__file__).parent / 'fixtures' / 'sample_translation.docx'
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    
    print(f"✅ Created sample Word file: {output_path}")
    return output_path


def test_word_parsing():
    """Test parsing Word document"""
    print("\n" + "="*60)
    print("TEST: Word Document Parsing")
    print("="*60)
    
    # Create sample file
    sample_path = create_sample_word_file()
    
    # Parse it
    parser = WordParser(str(sample_path))
    segments = parser.parse()
    
    print(f"\nFound {len(segments)} segments:")
    for i, seg in enumerate(segments, 1):
        print(f"  {i}. ID: {seg.id_text}")
        print(f"     Source: {seg.source_text[:50]}...")
    
    assert len(segments) == 5, f"Expected 5 segments, got {len(segments)}"
    assert segments[0].source_text == "Welcome to the course!"
    assert "%TotalQuestions%%" in segments[2].source_text
    
    print("\n✅ Parsing test passed!")
    return parser, segments


def test_word_reconstruction():
    """Test reconstructing Word document with translations"""
    print("\n" + "="*60)
    print("TEST: Word Document Reconstruction")
    print("="*60)
    
    parser, segments = test_word_parsing()
    
    # Add fake translations
    print("\nAdding translations...")
    for seg in segments:
        seg.target_text = f"[ES] {seg.source_text}"
    
    # Reconstruct
    output_path = Path(__file__).parent / 'fixtures' / 'translated_output.docx'
    success = parser.reconstruct(segments, str(output_path))
    
    assert success, "Reconstruction should succeed"
    print(f"✅ Saved translated file: {output_path}")
    
    # Verify by parsing the output
    print("\nVerifying output...")
    output_parser = WordParser(str(output_path))
    output_segments = output_parser.parse()
    
    assert len(output_segments) == len(segments)
    assert "[ES]" in output_segments[0].target_text
    
    print("✅ Verification passed!")
    
    # Cleanup
    output_path.unlink()
    print("✅ Cleaned up test files")


if __name__ == "__main__":
    print("="*60)
    print("Word Parser Tests")
    print("="*60)
    
    try:
        test_word_reconstruction()
        print("\n" + "="*60)
        print("🎉 ALL TESTS PASSED!")
        print("="*60)
    except AssertionError as e:
        print(f"\n❌ TEST FAILED: {e}")
        exit(1)
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        exit(1)