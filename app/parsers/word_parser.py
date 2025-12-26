"""
Word Document Parser for Articulate Storyline Translation
Handles .docx translation export format with MULTIPLE TABLES
"""

from docx import Document
from docx.shared import Pt, RGBColor
from typing import List, Dict, Tuple
from pathlib import Path


class WordSegment:
    """Represents a translatable segment from Word document"""
    
    def __init__(self, row_index: int, id_text: str, source_text: str, 
                 target_text: str = "", context: str = "", table_index: int = 0):
        self.row_index = row_index
        self.table_index = table_index  # NEW: Track which table this segment belongs to
        self.id_text = id_text
        self.source_text = source_text
        self.target_text = target_text
        self.translation = target_text
        self.context = context
        self.source_cell = None
        self.target_cell = None


class WordParser:
    """Parser for Word documents exported from Articulate Storyline"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self.segments: List[WordSegment] = []
        self.table_mappings = []  # Store column mappings for each table
        
    def parse(self) -> List[WordSegment]:
        """Parse Word document and extract segments from ALL tables"""
        try:
            self.doc = Document(self.file_path)
            
            # Find ALL translation tables
            tables = self._find_all_translation_tables()
            
            if not tables:
                raise ValueError("No translation tables found in document")
            
            print(f"Found {len(tables)} translation tables")
            
            # Extract segments from all tables
            all_segments = []
            for table_idx, table in enumerate(tables):
                # Identify column structure for this table
                col_mapping = self._identify_columns(table, table_idx)
                self.table_mappings.append(col_mapping)
                
                # Extract segments from this table
                segments = self._extract_segments(table, table_idx, col_mapping)
                all_segments.extend(segments)
                print(f"Table {table_idx}: Extracted {len(segments)} segments")
            
            self.segments = all_segments
            return self.segments
            
        except Exception as e:
            raise RuntimeError(f"Failed to parse Word document: {str(e)}")
    
    def _find_all_translation_tables(self):
        """Find ALL translation tables in the document"""
        translation_tables = []
        
        for table in self.doc.tables:
            if len(table.rows) > 1 and len(table.columns) >= 3:
                # Check if first row looks like headers
                first_row_text = ' '.join(cell.text.lower() for cell in table.rows[0].cells)
                
                # Look for translation table indicators
                if any(keyword in first_row_text for keyword in 
                       ['source text', 'translation', 'type', 'id']):
                    translation_tables.append(table)
        
        return translation_tables
    
    def _identify_columns(self, table, table_idx):
        """Identify which columns contain source, target, and ID for a specific table"""
        if len(table.rows) == 0:
            raise ValueError(f"Table {table_idx} has no rows")
        
        header_row = table.rows[0]
        col_mapping = {
            'id': None,
            'source': None,
            'target': None
        }
        
        # Check header text to identify columns
        for i, cell in enumerate(header_row.cells):
            cell_text = cell.text.lower().strip()
            
            # Identify ID column
            if 'id' in cell_text and '🔒' in cell_text:
                col_mapping['id'] = i
            
            # Identify source column
            elif 'source text' in cell_text:
                col_mapping['source'] = i
            
            # Identify target/translation column
            elif 'translation' in cell_text:
                col_mapping['target'] = i
        
        # Fallback: Use column positions if headers unclear
        if col_mapping['source'] is None:
            # Source is usually second-to-last
            col_mapping['source'] = max(0, len(header_row.cells) - 2)
        
        if col_mapping['target'] is None:
            # Target is usually last
            col_mapping['target'] = len(header_row.cells) - 1
        
        print(f"Table {table_idx} columns: ID={col_mapping['id']}, "
              f"Source={col_mapping['source']}, Target={col_mapping['target']}")
        
        return col_mapping
    
    def _extract_segments(self, table, table_idx, col_mapping) -> List[WordSegment]:
        """Extract text segments from a specific table"""
        segments = []
        
        # Skip header row (row 0)
        for row_idx, row in enumerate(table.rows[1:], start=1):
            try:
                cells = row.cells
                
                # Get ID if column exists
                id_text = ""
                if col_mapping['id'] is not None and col_mapping['id'] < len(cells):
                    id_text = cells[col_mapping['id']].text.strip()
                
                # Get source text
                if col_mapping['source'] >= len(cells):
                    continue
                source_cell = cells[col_mapping['source']]
                source_text = source_cell.text.strip()
                
                # Skip empty rows
                if not source_text:
                    continue
                
                # Get target text (may be empty)
                target_text = ""
                target_cell = None
                if col_mapping['target'] < len(cells):
                    target_cell = cells[col_mapping['target']]
                    target_text = target_cell.text.strip()
                
                # Create segment
                segment = WordSegment(
                    row_index=row_idx,
                    table_index=table_idx,  # Track which table
                    id_text=id_text or f"table{table_idx}_row{row_idx}",
                    source_text=source_text,
                    target_text=target_text,
                    context=f"Table {table_idx}, Row {row_idx}"
                )
                
                # Store cell references for reconstruction
                segment.source_cell = source_cell
                segment.target_cell = target_cell
                
                segments.append(segment)
                
            except Exception as e:
                print(f"Warning: Failed to parse table {table_idx}, row {row_idx}: {e}")
                continue
        
        return segments
    
    def reconstruct(self, translated_segments: List[WordSegment], 
                   output_path: str) -> bool:
        """Reconstruct Word document with translations in ALL tables"""
        try:
            if not self.doc:
                raise ValueError("No document loaded. Call parse() first.")
            
            # Find all tables again
            tables = self._find_all_translation_tables()
            
            if not tables:
                raise ValueError("Translation tables not found")
            
            # Group segments by table
            segments_by_table = {}
            for segment in translated_segments:
                table_idx = segment.table_index
                if table_idx not in segments_by_table:
                    segments_by_table[table_idx] = []
                segments_by_table[table_idx].append(segment)
            
            # Update each table
            for table_idx, table in enumerate(tables):
                if table_idx not in segments_by_table:
                    print(f"No translations for table {table_idx}")
                    continue
                
                col_mapping = self.table_mappings[table_idx]
                segments = segments_by_table[table_idx]
                
                print(f"Updating table {table_idx} with {len(segments)} translations")
                
                # Update each segment in this table
                for segment in segments:
                    if segment.row_index > len(table.rows) - 1:
                        print(f"Warning: Row {segment.row_index} out of range in table {table_idx}")
                        continue
                    
                    # Get the actual row
                    row = table.rows[segment.row_index]
                    
                    target_col = col_mapping['target']
                    if target_col >= len(row.cells):
                        print(f"Warning: Target column not found in table {table_idx}, row {segment.row_index}")
                        continue
                    
                    # Get target cell
                    target_cell = row.cells[target_col]
                    
                    # Get the translation
                    translation = getattr(segment, 'target_text', 
                                        getattr(segment, 'translation', ''))
                    
                    if not translation:
                        continue
                    
                    # Clear existing content
                    for paragraph in target_cell.paragraphs:
                        paragraph.clear()
                    
                    # Add translated text
                    if target_cell.paragraphs:
                        paragraph = target_cell.paragraphs[0]
                    else:
                        paragraph = target_cell.add_paragraph()
                    
                    paragraph.text = translation
                    
                    # Copy formatting from source cell
                    if segment.source_cell and segment.source_cell.paragraphs:
                        source_para = segment.source_cell.paragraphs[0]
                        if source_para.runs and paragraph.runs:
                            source_run = source_para.runs[0]
                            target_run = paragraph.runs[0]
                            if source_run.bold:
                                target_run.bold = source_run.bold
                            if source_run.italic:
                                target_run.italic = source_run.italic
            
            # Save document
            self.doc.save(output_path)
            print(f"Document saved to {output_path}")
            return True
            
        except Exception as e:
            print(f"Error reconstructing Word document: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    @staticmethod
    def is_word_translation_format(file_path: str) -> bool:
        """Check if a Word file appears to be a Storyline translation export"""
        try:
            doc = Document(file_path)
            
            if not doc.tables:
                return False
            
            # Check if any table looks like a translation table
            for table in doc.tables:
                if len(table.rows) < 2 or len(table.columns) < 2:
                    continue
                
                header_text = ' '.join(cell.text.lower() for cell in table.rows[0].cells)
                
                if any(keyword in header_text for keyword in 
                      ['translation', 'source text', 'type']):
                    return True
            
            return False
            
        except:
            return False


def parse_word(file_path: str) -> List[WordSegment]:
    """Convenience function to parse a Word document"""
    parser = WordParser(file_path)
    return parser.parse()